from __future__ import annotations

import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
ASSETS = OUT / "assets"
SYSTEM_DOCX = OUT / "sudan_ai_system_design_document.docx"
DECK_DOCX = OUT / "sudan_ai_product_deck.docx"

BRAND = "IBRAHEM AHMED HASSAN"
WEBSITE = "ibrahemahmed.com"
EMAIL = "hello@brahemahmed.com"
LINKEDIN = "https://www.linkedin.com/in/ibrahem-ahmed-hassan/"

INK = RGBColor(8, 34, 49)
MUTED = RGBColor(89, 104, 112)
TEAL = RGBColor(0, 126, 109)
GREEN = RGBColor(15, 137, 83)
BLUE = RGBColor(20, 78, 124)
LIGHT = "F2F7F6"
LIGHT_BLUE = "EAF2F7"
LIGHT_GREEN = "EAF7F1"
BORDER = "D8E2E1"


def font_path(name: str) -> str:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("/usr/share/fonts/truetype/dejavu") / name,
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return ""


FONT_REG = font_path("arial.ttf") or font_path("DejaVuSans.ttf")
FONT_BOLD = font_path("arialbd.ttf") or font_path("DejaVuSans-Bold.ttf")


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    path = FONT_BOLD if bold else FONT_REG
    if path:
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if width <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_card(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str,
    fill: str,
    outline: str = "#D8E2E1",
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline, width=2)
    title_font = pil_font(34, True)
    body_font = pil_font(23)
    draw.text((x1 + 28, y1 + 24), title, font=title_font, fill="#082231")
    y = y1 + 76
    for line in wrap_text(draw, body, body_font, x2 - x1 - 56):
        draw.text((x1 + 28, y), line, font=body_font, fill="#4F636D")
        y += 34


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#007E6D") -> None:
    draw.line([start, end], fill=color, width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for offset in (math.pi * 0.82, -math.pi * 0.82):
        point = (
            end[0] - length * math.cos(angle + offset),
            end[1] - length * math.sin(angle + offset),
        )
        draw.line([end, point], fill=color, width=5)


def make_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1050), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = pil_font(44, True)
    draw.text((70, 48), "Production Architecture", font=title_font, fill="#082231")
    draw.text(
        (70, 106),
        "Arabic-first RAG learning platform for Sudan middle school textbooks",
        font=pil_font(25),
        fill="#596870",
    )

    cards = [
        ((70, 190, 430, 390), "Student UI", "Next.js 15, RTL Arabic, Tailwind, shadcn UI, Framer Motion"),
        ((550, 190, 910, 390), "API Layer", "Next.js route handlers for auth, books, chat, generation, admin"),
        ((1030, 190, 1390, 390), "AI Services", "OpenRouter chat plus Gemini Embedding 2 through provider abstraction"),
        ((550, 500, 910, 720), "RAG Core", "Selected-book retrieval, hybrid vector/text ranking, grounded prompt assembly"),
        ((1030, 500, 1390, 720), "Supabase", "Auth, PostgreSQL, pgvector, Storage, RLS policies, RPC retrieval"),
        ((70, 500, 430, 720), "Ingestion", "PDF upload/sync, text extraction, cleaning, chunking, embeddings"),
    ]
    fills = ["#EAF7F1", "#EAF2F7", "#F2F7F6", "#EAF7F1", "#EAF2F7", "#F7FAF9"]
    for (xy, title, body), fill in zip(cards, fills):
        draw_card(draw, xy, title, body, fill)

    arrow(draw, (430, 290), (550, 290))
    arrow(draw, (910, 290), (1030, 290))
    arrow(draw, (730, 390), (730, 500))
    arrow(draw, (910, 610), (1030, 610))
    arrow(draw, (430, 610), (550, 610))
    arrow(draw, (1210, 500), (1210, 390))

    draw.rounded_rectangle((70, 840, 1390, 970), radius=22, fill="#082231")
    draw.text((104, 868), "Deployment target", font=pil_font(31, True), fill="#FFFFFF")
    draw.text(
        (104, 916),
        "Vercel hosts the Next.js frontend and API routes. Supabase hosts Auth, Postgres, Storage, and pgvector.",
        font=pil_font(25),
        fill="#D9F3EF",
    )
    img.save(path)


def make_rag_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 50), "RAG Query and Ingestion Flow", font=pil_font(44, True), fill="#082231")
    steps = [
        ("1. PDF Intake", "Scan / upload textbooks, validate PDFs, record ingestion job."),
        ("2. Extract + Clean", "Read pages, normalize text, clean OCR and layout artifacts."),
        ("3. Chunk + Embed", "Chunk size 1200, overlap 200, embedding dimension 768."),
        ("4. Store Vectors", "Supabase documents table with pgvector and page metadata."),
        ("5. Ask Question", "Student selects book IDs, then sends a question."),
        ("6. Ground Answer", "Hybrid retrieval -> context -> OpenRouter -> Arabic response with citations."),
    ]
    x = 70
    y = 170
    w = 260
    h = 190
    for i, (title, body) in enumerate(steps):
        row = 0 if i < 3 else 1
        col = i if i < 3 else i - 3
        x = 70 + col * 555
        y = 170 + row * 310
        fill = "#EAF7F1" if i % 2 == 0 else "#EAF2F7"
        draw_card(draw, (x, y, x + 430, y + h), title, body, fill)
        if col < 2:
            arrow(draw, (x + 430, y + 95), (x + 515, y + 95))
    arrow(draw, (1185, 360), (1185, 480))
    draw.rounded_rectangle((70, 785, 1595, 850), radius=18, fill="#F7FAF9", outline="#D8E2E1")
    draw.text(
        (100, 803),
        "Grounding rule: every chat and generator request filters retrieval to selected_book_ids before calling the model.",
        font=pil_font(25, True),
        fill="#007E6D",
    )
    img.save(path)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths: list[int], indent: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def add_colored_rule(paragraph, color: str = "007E6D", size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    border = p_pr.find(qn("w:pBdr"))
    if border is None:
        border = OxmlElement("w:pBdr")
        p_pr.append(border)
    bottom = border.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        border.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)


def style_document(doc: Document, title: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    section.header.paragraphs[0].text = title
    section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(section.header.paragraphs[0].runs[0], size=8.5, color=MUTED)
    section.footer.paragraphs[0].text = f"Made by {BRAND} | {WEBSITE} | {EMAIL}"
    set_run_font(section.footer.paragraphs[0].runs[0], size=8.5, color=MUTED)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    for style_name, size, color, before, after in [
        ("Heading 1", 16, TEAL, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, INK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_page(doc: Document, title: str, subtitle: str, doc_type: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2)
    r = p.add_run(doc_type.upper())
    set_run_font(r, size=11, bold=True, color=TEAL)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)
    r = p.add_run(title)
    set_run_font(r, size=27, bold=True, color=INK)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=16)
    r = p.add_run(subtitle)
    set_run_font(r, size=13.5, color=MUTED)

    rule = doc.add_paragraph()
    add_colored_rule(rule)
    set_paragraph_spacing(rule, after=14)

    meta = doc.add_table(rows=5, cols=2)
    set_table_width(meta, [1800, 7200])
    rows = [
        ("Prepared by", BRAND),
        ("Website", WEBSITE),
        ("Email", EMAIL),
        ("LinkedIn", LINKEDIN),
        ("Document date", "May 30, 2026"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for idx, cell in enumerate(row.cells):
            set_cell_border(cell)
            set_cell_shading(cell, LIGHT if idx == 0 else "FFFFFF")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=0, line=1.1)
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(idx == 0), color=INK if idx == 0 else MUTED)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=18, after=6)
    r = p.add_run("Purpose")
    set_run_font(r, size=12, bold=True, color=TEAL)
    p = doc.add_paragraph(
        "This package documents the current production architecture, product direction, data management model, "
        "and scaling plan for the Sudan middle school AI learning platform."
    )
    set_paragraph_spacing(p, after=0)
    doc.add_page_break()


def add_h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4, line=1.15)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            set_paragraph_spacing(p, after=0)
            for r in p.runs:
                set_run_font(r, size=9.5, bold=True, color=INK)
    for row_values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, row_values):
            cell.text = value
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                set_paragraph_spacing(p, after=0, line=1.08)
                for r in p.runs:
                    set_run_font(r, size=9.2, color=INK)
    doc.add_paragraph()


def build_system_design() -> None:
    doc = Document()
    style_document(doc, "Sudan AI Learning Platform - System Design", "System design")
    doc.core_properties.author = BRAND
    doc.core_properties.title = "Sudan AI Learning Platform - System Design Document"

    add_title_page(
        doc,
        "Sudan Middle School AI Learning Platform",
        "System Design Document for a production-ready Arabic-first RAG learning product",
        "System Design Document",
    )

    add_h(doc, "1. Executive Summary")
    add_p(
        doc,
        "The platform is an Arabic-first AI learning product for Sudan middle school students. Students choose a grade, "
        "select one or multiple official textbooks, then ask questions or generate study material grounded only in those selected books."
    )
    add_bullets(
        doc,
        [
            "Current verified corpus: 46 total books in the database, including 30 official MDL intermediate books.",
            "Official book coverage: grade1 has 9 books, grade2 has 13 books, and grade3 has 8 books.",
            "Vector corpus: 6,292 total chunks, including 4,780 official-curriculum chunks.",
            "Core AI boundary: retrieval is filtered by selected_book_ids before the model receives context.",
        ]
    )

    add_h(doc, "2. Product Scope")
    add_table(
        doc,
        ["Area", "Implemented Capability", "Product Value"],
        [
            ["Auth", "Supabase email/password auth with protected student/admin routes.", "Keeps saved chats, flashcards, quizzes, and profile data private per user."],
            ["Library", "Grade filters, book search, PDF links, multi-book selection.", "Lets students study by class and subject without navigating a complex LMS."],
            ["AI Chat", "Streaming RAG chat over selected books only.", "Gives simple Arabic explanations while reducing hallucination risk."],
            ["Study Tools", "Flashcards, MCQ, exams, summaries, key points, study notes, Q/A drills.", "Turns textbooks into active-recall practice."],
            ["Dashboard", "Progress, saved flashcards, quiz scores, previous chats.", "Keeps study history useful beyond one session."],
            ["Admin", "Upload books, reprocess embeddings, inspect ingestion jobs, manage users and roles.", "Supports ongoing curriculum operations."],
        ],
        [1400, 4100, 3860],
    )

    add_h(doc, "3. High-Level Architecture")
    add_p(
        doc,
        "The architecture keeps the product simple to operate while preserving clear boundaries between UI, API, retrieval, AI providers, storage, and data ownership."
    )
    doc.add_picture(str(ASSETS / "architecture.png"), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_h(doc, "4. Runtime Components")
    add_table(
        doc,
        ["Component", "Technology", "Responsibility"],
        [
            ["Frontend", "Next.js 15, TypeScript, TailwindCSS, shadcn-style primitives, Framer Motion", "Arabic RTL application shell, landing page, library, study workspace, dashboard, history, admin."],
            ["API routes", "Next.js route handlers", "Books API, RAG chat, study generation, chat history, quiz score, admin upload/reprocess/users."],
            ["Auth", "Supabase Auth + middleware", "Protects dashboard, library, study, history, flashcard export, and admin routes."],
            ["Database", "Supabase PostgreSQL", "Stores books, documents, profiles, chats, flashcards, quizzes, bookmarks, progress, and ingestion jobs."],
            ["Vector DB", "Supabase pgvector", "Stores 768-dimensional embeddings and supports HNSW vector search."],
            ["AI", "OpenRouter chat plus provider abstraction", "Allows switching among DeepSeek, Claude, GPT, Gemini, Qwen, or other supported chat models."],
            ["Embeddings", "OpenRouter Gemini Embedding 2 Preview with Gemini/Jina fallback paths", "Embeds document chunks and query text for retrieval."],
            ["Storage", "Supabase Storage bucket books", "Stores validated PDF textbooks and public PDF URLs."],
        ],
        [1450, 3000, 4910],
    )

    add_h(doc, "5. RAG Pipeline")
    doc.add_picture(str(ASSETS / "rag_pipeline.png"), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bullets(
        doc,
        [
            "Ingestion scans PDFs, detects title/grade/subject where possible, uploads the PDF, extracts pages, cleans text, and chunks with size 1200 and overlap 200.",
            "Embedding is batched and stored in documents.embedding as vector(768), with metadata for book, subject, grade, page, source file, and official-curriculum status.",
            "At query time the API embeds the student question, calls match_documents with selected_book_ids, builds a cited context block, and sends only that context to OpenRouter.",
            "The system prompt tells the assistant to answer only from context and to say the information is not in the selected book when context is missing.",
        ]
    )

    add_h(doc, "6. Database and Data Ownership")
    add_table(
        doc,
        ["Table", "Primary Owner", "Purpose"],
        [
            ["classes, subjects", "Admin/public read", "Curriculum taxonomy and grade navigation."],
            ["books, chapters", "Admin writes, students read", "Book catalog, subject, grade, cover/PDF metadata."],
            ["documents", "Admin/service writes, authenticated read through retrieval", "Chunk text, vector embedding, page number, metadata, search vector."],
            ["profiles", "User/admin", "Student name, grade, and admin role."],
            ["chat_sessions, chat_messages", "User", "Saved per-user chat history with selected book IDs and citations."],
            ["flashcards", "User", "Saved active-recall cards generated from selected books."],
            ["quizzes, quiz_questions", "User", "MCQ/exam records, score, explanations, difficulty."],
            ["bookmarks, study_progress", "User", "Lesson notes, progress, completed lessons."],
            ["ingestion_jobs", "Admin", "Operational status for uploads and reprocessing."],
        ],
        [2100, 2100, 5160],
    )

    add_h(doc, "7. Security and Authorization")
    add_bullets(
        doc,
        [
            "Supabase Row Level Security is enabled across application tables.",
            "Students can manage only their own flashcards, quizzes, progress, bookmarks, chat sessions, and chat messages.",
            "Admin endpoints use route-level admin guards and service-role operations only where required for ingestion or user management.",
            "The public books API can expose official books, while scope=all is admin-only.",
            "AI requests are rate limited per user to control abuse and cost.",
            "Uploads validate MIME type, file size, and PDF requirement before storage or processing.",
            "Secrets stay in environment variables: Supabase service key, OpenRouter API key, embedding provider keys, and admin emails.",
        ]
    )

    add_h(doc, "8. Data Management Model")
    add_p(
        doc,
        "The data strategy should treat textbooks as versioned curriculum assets, not one-off files. Every PDF should carry structured metadata, stable naming, ingest status, and retrievable text quality signals."
    )
    add_table(
        doc,
        ["Concern", "Recommended Practice", "Reason"],
        [
            ["Book identity", "Use grade, subject, normalized title, source URL, curriculum year/version.", "Prevents duplicate or ambiguous books when more stages are added."],
            ["Text quality", "Store extraction method, OCR status, page count, chunk count, failed pages.", "Makes scanned/image PDFs visible to admins instead of silent failures."],
            ["Reprocessing", "Keep ingestion_jobs and allow forced reprocess by book/version.", "Lets admins improve OCR, embeddings, or chunking without losing the catalog."],
            ["Metadata", "Standardize subject slugs and grade/stage IDs.", "Supports all classes and subjects later without UI rewrites."],
            ["Backups", "Export Supabase schema, storage manifest, and ingestion logs regularly.", "Protects curriculum data and supports migration."],
        ],
        [1750, 3820, 3790],
    )

    add_h(doc, "9. Scaling Plan")
    add_table(
        doc,
        ["Layer", "Next Step", "Scale Benefit"],
        [
            ["Curriculum", "Generalize classes into stages, grades, subjects, terms, and curriculum versions.", "Allows primary, secondary, and future curricula without special-case UI."],
            ["Ingestion", "Move PDF processing to background workers with queue retries.", "Keeps admin uploads responsive and prevents Vercel function timeouts."],
            ["OCR", "Add Arabic OCR pipeline for scanned books using Tesseract, PaddleOCR, Google Document AI, or Azure Document Intelligence.", "Unlocks image-only books that currently have no retrievable text chunks."],
            ["Retrieval", "Partition or filter indexes by grade/subject/book and tune HNSW parameters.", "Maintains low-latency search as vectors grow from thousands to millions."],
            ["Caching", "Cache embeddings for repeated queries and cache common generation outputs per book/topic.", "Reduces OpenRouter/Gemini spend and improves perceived speed."],
            ["Observability", "Add structured logs for retrieval score, model, latency, token cost, and refusal rate.", "Supports quality tuning and cost control."],
            ["Product", "Add teacher/admin analytics, premium exports, PWA/offline reading, and school/team accounts.", "Creates stronger retention and monetization paths."],
        ],
        [1600, 4200, 3560],
    )

    add_h(doc, "10. Future Roadmap")
    add_h(doc, "10.1 Add Other Subjects and Classes", level=2)
    add_bullets(
        doc,
        [
            "Expand the taxonomy from grade1/grade2/grade3 to stage -> grade -> term -> subject -> book.",
            "Add all remaining subjects for each class, then add primary and secondary school stages using the same ingestion pipeline.",
            "Build an admin mapping screen for subject aliases, curriculum source URLs, and version labels.",
            "Add curriculum completeness dashboards: expected books vs uploaded books vs chunked books vs OCR-needed books.",
        ]
    )
    add_h(doc, "10.2 Product Enhancements", level=2)
    add_bullets(
        doc,
        [
            "Teacher mode: assign practice exams, view class results, and publish recommended study sets.",
            "Premium mode: advanced exam builder, printable study packs, PDF exports, analytics, and priority AI quota.",
            "Mobile PWA: offline saved flashcards, recent chats, and low-bandwidth mode for students.",
            "Quality controls: answer citations, source-page viewer, user feedback, and admin review of weak answers.",
        ]
    )
    add_h(doc, "10.3 Engineering Enhancements", level=2)
    add_bullets(
        doc,
        [
            "Add background job orchestration for ingestion, OCR, embedding, and reprocessing.",
            "Create automated E2E tests for signup, library selection, RAG chat, flashcards, MCQ scoring, history, and admin upload.",
            "Add model routing so easy questions use cheaper models and exam generation can use stronger models.",
            "Add analytics events for book opens, chat success, generation type, quiz completion, and retention.",
        ]
    )

    add_h(doc, "11. Deployment and Operations")
    add_bullets(
        doc,
        [
            "Frontend and API: deploy Next.js to Vercel with environment variables for Supabase and OpenRouter.",
            "Backend services: Supabase hosts Auth, PostgreSQL, pgvector, Storage, RLS, and SQL RPC functions.",
            "Operational commands: npm run build, npm run lint, npm run typecheck, npm run sync:mdl, npm run ingest:mdl, npm run ingest:books.",
            "Runbooks should include key rotation, failed ingestion recovery, OCR reprocessing, and monthly cost review.",
        ]
    )

    add_h(doc, "12. Key Risks and Mitigations")
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Scanned PDFs with no extractable text", "AI cannot answer from those books.", "Add OCR and mark OCR-needed books in admin."],
            ["Large vector growth", "Retrieval latency and storage cost increase.", "Use metadata filters, index tuning, partitioning, and archival policy."],
            ["Model hallucination", "Student trust and academic accuracy risk.", "Keep selected-book filtering, refusal prompt, citations, and answer feedback loop."],
            ["API key leakage", "Cost and security exposure.", "Keep keys server-side, rotate regularly, never expose service role to the client."],
            ["Curriculum changes", "Old answers may mismatch new editions.", "Version books and associate chunks with source/year/version metadata."],
        ],
        [2450, 2850, 4060],
    )

    doc.save(SYSTEM_DOCX)


def set_deck_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.48)
    section.right_margin = Inches(0.48)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    doc.core_properties.author = BRAND
    doc.core_properties.title = "Sudan AI Learning Platform - Product Deck"
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(13)
    normal.font.color.rgb = INK


def add_slide_number_footer(doc: Document, number: int) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"{number:02d} | Made by {BRAND} | {WEBSITE}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=8.5, color=MUTED)


def deck_title(doc: Document, kicker: str, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2)
    r = p.add_run(kicker.upper())
    set_run_font(r, size=11, bold=True, color=TEAL)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4)
    r = p.add_run(title)
    set_run_font(r, size=28, bold=True, color=INK)
    if subtitle:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=8)
        r = p.add_run(subtitle)
        set_run_font(r, size=13, color=MUTED)


def slide_cards(doc: Document, cards: list[tuple[str, str]], cols: int = 3) -> None:
    rows = math.ceil(len(cards) / cols)
    table = doc.add_table(rows=rows, cols=cols)
    set_table_width(table, [int(9360 / cols)] * cols, indent=0)
    for idx in range(rows * cols):
        cell = table.cell(idx // cols, idx % cols)
        set_cell_border(cell, "D8E2E1")
        set_cell_shading(cell, "FFFFFF" if idx % 2 else LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        if idx < len(cards):
            title, body = cards[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=3)
            r = p.add_run(title)
            set_run_font(r, size=14, bold=True, color=TEAL)
            p = cell.add_paragraph()
            set_paragraph_spacing(p, after=0, line=1.08)
            r = p.add_run(body)
            set_run_font(r, size=10.7, color=INK)
        else:
            cell.text = ""
    doc.add_paragraph()


def add_slide_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_product_deck() -> None:
    doc = Document()
    set_deck_layout(doc)

    # Slide 1
    add_slide_number_footer(doc, 1)
    cover = doc.add_table(rows=1, cols=2)
    set_table_width(cover, [4700, 4660], indent=0)
    left, right = cover.rows[0].cells
    for cell in (left, right):
        set_cell_border(cell, "FFFFFF", "0")
    p = left.paragraphs[0]
    set_paragraph_spacing(p, after=5)
    r = p.add_run("PRODUCT DECK")
    set_run_font(r, size=11, bold=True, color=TEAL)
    p = left.add_paragraph()
    set_paragraph_spacing(p, after=8)
    r = p.add_run("Sudan Middle School AI Learning Platform")
    set_run_font(r, size=29, bold=True, color=INK)
    p = left.add_paragraph()
    set_paragraph_spacing(p, after=18, line=1.15)
    r = p.add_run("A grounded Arabic AI study assistant built from Sudanese curriculum books.")
    set_run_font(r, size=15, color=MUTED)
    for text in [
        f"Made by {BRAND}",
        WEBSITE,
        EMAIL,
        LINKEDIN,
    ]:
        p = left.add_paragraph()
        set_paragraph_spacing(p, after=3)
        r = p.add_run(text)
        set_run_font(r, size=11.5, color=INK if text.startswith("Made") else MUTED, bold=text.startswith("Made"))
    hero = ROOT / "public" / "images" / "landing-hero.png"
    if hero.exists():
        right.paragraphs[0].add_run().add_picture(str(hero), width=Inches(4.3))
    add_slide_break(doc)

    # Slide 2
    add_slide_number_footer(doc, 2)
    deck_title(doc, "Problem", "Students need trusted help from the actual textbook")
    slide_cards(
        doc,
        [
            ("Fragmented materials", "Students move between PDFs, notes, search results, and informal explanations."),
            ("Hallucination risk", "Generic AI can answer outside the Sudan curriculum and confuse students."),
            ("Low active recall", "Students need flashcards, MCQs, and exams, not only passive summaries."),
            ("Admin overhead", "Schools and operators need a repeatable way to add books and keep data clean."),
        ],
        cols=4,
    )
    add_slide_break(doc)

    # Slide 3
    add_slide_number_footer(doc, 3)
    deck_title(doc, "Solution", "A ChatPDF plus Quizlet experience for Sudan middle school")
    slide_cards(
        doc,
        [
            ("Choose class", "First, second, or third intermediate grade."),
            ("Pick books", "Select one book or multiple books at the same time."),
            ("Ask AI", "The assistant answers only from selected textbook chunks."),
            ("Generate study tools", "Flashcards, MCQ, exams, summaries, key points, notes, and drills."),
            ("Save progress", "Previous chats, cards, quizzes, scores, and bookmarks stay attached to the user."),
            ("Admin controls", "Upload, rename, reprocess, and inspect ingestion status."),
        ],
        cols=3,
    )
    add_slide_break(doc)

    # Slide 4
    add_slide_number_footer(doc, 4)
    deck_title(doc, "Current Build", "Production-ready architecture already implemented")
    add_table(
        doc,
        ["Metric", "Current State"],
        [
            ["Official books", "30 MDL intermediate books"],
            ["Grade distribution", "grade1: 9, grade2: 13, grade3: 8"],
            ["Vector corpus", "6,292 total chunks, 4,780 official-curriculum chunks"],
            ["Core routes", "dashboard, library, study, history, admin, flashcards export, landing page"],
            ["AI providers", "OpenRouter chat, OpenRouter Gemini Embedding 2 Preview, Gemini/Jina fallback paths"],
        ],
        [3000, 6360],
    )
    add_slide_break(doc)

    # Slide 5
    add_slide_number_footer(doc, 5)
    deck_title(doc, "How RAG Works", "Selected-book filtering is the product moat")
    doc.add_picture(str(ASSETS / "rag_pipeline.png"), width=Inches(9.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_slide_break(doc)

    # Slide 6
    add_slide_number_footer(doc, 6)
    deck_title(doc, "Student Experience", "Simple Arabic workflows for daily study")
    slide_cards(
        doc,
        [
            ("Explain lessons", "Ask for step-by-step Arabic explanation from the selected book."),
            ("Summarize chapters", "Turn long textbook sections into concise study summaries."),
            ("Practice recall", "Generate flashcards and Q/A drills for quick revision."),
            ("Test readiness", "Create MCQs and timed practice exams with scoring and explanations."),
        ],
        cols=4,
    )
    add_slide_break(doc)

    # Slide 7
    add_slide_number_footer(doc, 7)
    deck_title(doc, "Admin and Data Operations", "A repeatable path to manage curriculum data")
    slide_cards(
        doc,
        [
            ("Upload PDFs", "Validate PDF type and size, store in Supabase Storage."),
            ("Metadata control", "Edit book title, grade, subject, source file, and official scope."),
            ("Reprocess embeddings", "Regenerate chunks and vectors when extraction, OCR, or metadata improves."),
            ("Ingestion status", "Track queued, running, completed, and failed jobs for operations."),
            ("Data quality", "Monitor chunk count, OCR-needed PDFs, failed pages, and duplicate titles."),
            ("Role management", "Use admin-only routes and RLS to separate operators from students."),
        ],
        cols=3,
    )
    add_slide_break(doc)

    # Slide 8
    add_slide_number_footer(doc, 8)
    deck_title(doc, "Expansion Plan", "Add every subject, every class, and later every stage")
    slide_cards(
        doc,
        [
            ("Taxonomy upgrade", "Move from grade IDs to stage -> grade -> term -> subject -> book version."),
            ("Curriculum completeness", "Show expected vs uploaded vs processed vs OCR-needed books."),
            ("OCR layer", "Add Arabic OCR for scanned PDFs so image-only books become searchable."),
            ("Teacher mode", "Assignments, class analytics, recommended cards, and exam publishing."),
            ("Premium exports", "Printable PDF study packs, advanced exams, and priority AI quota."),
            ("PWA/mobile", "Low-bandwidth mobile study, offline flashcards, and saved recent chats."),
        ],
        cols=3,
    )
    add_slide_break(doc)

    # Slide 9
    add_slide_number_footer(doc, 9)
    deck_title(doc, "Scaling Strategy", "From one curriculum stage to a national learning graph")
    add_table(
        doc,
        ["Scale Area", "Plan"],
        [
            ["Vectors", "Use metadata filters, HNSW tuning, and partitioning by stage/grade/subject as corpus grows."],
            ["Ingestion", "Move heavy PDF, OCR, and embedding work to background queues with retry and audit logs."],
            ["AI cost", "Cache repeated embeddings, route easy tasks to cheaper models, and track token spend per feature."],
            ["Quality", "Store answer feedback, retrieval scores, refusal rates, and reviewed weak-answer examples."],
            ["Data governance", "Version every curriculum edition and keep source URL, file hash, OCR status, and chunk count."],
        ],
        [2800, 6560],
    )
    add_slide_break(doc)

    # Slide 10
    add_slide_number_footer(doc, 10)
    deck_title(doc, "Next 90 Days", "Turn the platform into a stronger education product")
    slide_cards(
        doc,
        [
            ("1. Complete OCR", "Process scanned/image PDFs and expose OCR status in admin."),
            ("2. Expand curriculum", "Add all remaining subjects and prepare primary/secondary stage taxonomy."),
            ("3. Improve learning loops", "Teacher mode, assignments, analytics, and personalized revision."),
            ("4. Monetize carefully", "Premium study packs, advanced exams, school accounts, and priority AI usage."),
        ],
        cols=4,
    )
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14, after=0)
    r = p.add_run(f"Contact: {BRAND} | {WEBSITE} | {EMAIL} | {LINKEDIN}")
    set_run_font(r, size=13, bold=True, color=TEAL)

    doc.save(DECK_DOCX)


def main() -> None:
    OUT.mkdir(exist_ok=True)
    ASSETS.mkdir(exist_ok=True)
    make_architecture_diagram(ASSETS / "architecture.png")
    make_rag_diagram(ASSETS / "rag_pipeline.png")
    build_system_design()
    build_product_deck()
    print(SYSTEM_DOCX)
    print(DECK_DOCX)


if __name__ == "__main__":
    main()
